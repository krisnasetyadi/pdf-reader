import pdfplumber
import os
import logging
import re
import subprocess
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS # type: ignore
from langchain_core.documents import Document
from config import config

logger = logging.getLogger(__name__)


def extract_text_from_pdf(pdf_path):
    """Extract page-by-page clean text and page numbers from PDF"""
    pages_data = []
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    # Clean text
                    page_text = re.sub(r'\s+', ' ', page_text).strip()
                    if page_text:
                        pages_data.append({
                            "text": page_text,
                            "page": page_num
                        })
        return pages_data if pages_data else None
    except Exception as e:
        logger.error(f"Failed to extract text from {pdf_path}: {str(e)}")
        return None


def extract_text_from_docx(docx_path):
    """Extract paragraph + table text from a Word document as one page-like
    block — same {"text", "page"} shape as extract_text_from_pdf so the
    downstream chunking/indexing code doesn't need to know the source format."""
    try:
        import docx
        document = docx.Document(docx_path)
        parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)
        text = re.sub(r'\s+', ' ', "\n".join(parts)).strip()
        return [{"text": text, "page": 1}] if text else None
    except Exception as e:
        logger.error(f"Failed to extract text from {docx_path}: {str(e)}")
        return None


def extract_text_from_doc(doc_path):
    """Extract text from a legacy binary .doc file via the `antiword` CLI.
    python-docx only understands the modern .docx XML format — the old
    binary .doc layout has no reliable pure-Python parser, so this shells
    out like the app already does elsewhere (e.g. reportlab for PDF export).
    Requires the `antiword` system package (added to the Dockerfile for
    deployment); if it's missing — e.g. a local dev machine without it —
    this logs and returns None, same as every other extractor's failure path,
    instead of crashing the whole upload."""
    try:
        result = subprocess.run(
            [config.antiword_path, doc_path],
            capture_output=True,
            text=True,
            timeout=30,
        )
        if result.returncode != 0:
            logger.error(f"antiword failed for {doc_path}: {result.stderr.strip()}")
            return None
        text = re.sub(r'\s+', ' ', result.stdout).strip()
        return [{"text": text, "page": 1}] if text else None
    except FileNotFoundError:
        logger.error(
            f"antiword not found at '{config.antiword_path}' — cannot extract text from "
            "legacy .doc files. Install it (e.g. `apt-get install antiword`), or if it's "
            "already installed but not on this process's PATH, set ANTIWORD_PATH in .env "
            "to its full path."
        )
        return None
    except Exception as e:
        logger.error(f"Failed to extract text from {doc_path}: {str(e)}")
        return None


def extract_text_from_csv(csv_path):
    """Extract CSV rows as readable "header: value" text — one page-like block."""
    try:
        import csv as csv_module
        rows_text = []
        with open(csv_path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv_module.reader(f)
            header = next(reader, None)
            for row in reader:
                if header and len(header) == len(row):
                    line = ", ".join(f"{h}: {v}" for h, v in zip(header, row))
                else:
                    line = ", ".join(row)
                if line.strip(", "):
                    rows_text.append(line)
        text = "\n".join(rows_text).strip()
        return [{"text": text, "page": 1}] if text else None
    except Exception as e:
        logger.error(f"Failed to extract text from {csv_path}: {str(e)}")
        return None


def extract_text_from_xlsx(xlsx_path):
    """Extract every sheet's rows as readable text — one page-like block per sheet
    (so a multi-sheet workbook chunks/cites per sheet, similar to per-page PDF text)."""
    try:
        import openpyxl
        workbook = openpyxl.load_workbook(xlsx_path, data_only=True, read_only=True)
        pages_data = []
        for sheet_index, sheet in enumerate(workbook.worksheets, 1):
            rows_iter = sheet.iter_rows(values_only=True)
            header = next(rows_iter, None)
            lines = []
            for row in rows_iter:
                if header and len(header) == len(row):
                    line = ", ".join(f"{h}: {v}" for h, v in zip(header, row) if v is not None)
                else:
                    line = ", ".join(str(v) for v in row if v is not None)
                if line.strip(", "):
                    lines.append(line)
            sheet_text = "\n".join(lines).strip()
            if sheet_text:
                pages_data.append({"text": f"[Sheet: {sheet.title}]\n{sheet_text}", "page": sheet_index})
        return pages_data if pages_data else None
    except Exception as e:
        logger.error(f"Failed to extract text from {xlsx_path}: {str(e)}")
        return None


def extract_text_from_txt(txt_path):
    """Read a plain-text file, trying a few encodings — one page-like block."""
    text = None
    for enc in ("utf-8", "utf-8-sig", "utf-16", "latin-1"):
        try:
            with open(txt_path, "r", encoding=enc) as f:
                text = f.read()
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        logger.error(f"Failed to decode text file {txt_path} with any supported encoding")
        return None
    text = text.strip()
    return [{"text": text, "page": 1}] if text else None


# Dispatch table for process_pdfs — keyed by lowercase file extension. PDF
# stays the default fallback so any caller relying on the old PDF-only
# behavior (e.g. a file with no/unknown extension) sees no change.
DOCUMENT_EXTRACTORS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".doc": extract_text_from_doc,
    ".csv": extract_text_from_csv,
    ".xlsx": extract_text_from_xlsx,
    ".txt": extract_text_from_txt,
}

# Mirrors DOCUMENT_EXTRACTORS' key set — keyed by lowercase file extension,
# used wherever an uploaded document's real MIME type needs to be recorded
# (e.g. S3 object metadata) instead of assuming PDF.
CONTENT_TYPE_BY_EXT = {
    ".pdf": "application/pdf",
    ".txt": "text/plain",
    ".csv": "text/csv",
    ".doc": "application/msword",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


def process_pdfs(pdf_paths, collection_id):
    """Process uploaded documents (PDF, DOCX, CSV, XLSX, TXT — despite the
    name, kept for backward compatibility with existing callers) with
    improved page-level text splitting."""
    documents = []

    for pdf_path in pdf_paths:
        ext = os.path.splitext(pdf_path)[1].lower()
        extract_text = DOCUMENT_EXTRACTORS.get(ext, extract_text_from_pdf)
        pages_data = extract_text(pdf_path)
        if pages_data:
            for page_entry in pages_data:
                doc = Document(
                    page_content=page_entry["text"],
                    metadata={
                        "source": os.path.basename(pdf_path),
                        "file_path": pdf_path,
                        "collection_id": collection_id,
                        "page": str(page_entry["page"])
                    }
                )
                documents.append(doc)

    if not documents:
        logger.error("No text could be extracted from any PDF")
        return 0

    # Improved text splitting
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=config.chunk_size,
        chunk_overlap=config.chunk_overlap,
        separators=["\n\n", "\n", ". ", "! ", "? ", "; ", ", ", " ", ""],
        length_function=len,
        keep_separator=True
    )

    chunks = text_splitter.split_documents(documents)

    if not chunks:
        logger.error("No chunks created from documents")
        return 0

    # Create vector store
    try:
        from processor import processor
        vector_store = FAISS.from_documents(chunks, processor.embeddings)

        # Save vector store
        index_path = os.path.join(config.index_folder, collection_id)
        os.makedirs(index_path, exist_ok=True)
        vector_store.save_local(index_path)

        logger.info(f"Created vector store with {len(chunks)} chunks for collection {collection_id}")
        return len(chunks)

    except Exception as e:
        logger.error(f"Failed to create vector store: {str(e)}")
        raise
